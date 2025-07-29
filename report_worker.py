import sys
import openpyxl as xl
import re
import math
import shutil
import threading
import time
import logging
from pathlib import Path
from collections import namedtuple, Counter
from docxtpl import DocxTemplate, InlineImage
from docx import Document
from docx.shared import Mm
from docxcompose.composer import Composer
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import win32com.client as win32

########################################
# 将日志信息输出到采用queue的Logger中
# 此Logger在主函数中定义，名字为“report”
########################################
log_show = logging.getLogger("report")


########################################
# 我国科学技术委员会正式颁布的《数字修约规则》，通常称为“四舍六入五成双”法则,即四舍六入五考虑。
# 当[尾数]≤4时舍去，尾数为6时进位。当尾数为5时，则应看末位数是奇数还是偶数，5前为偶数应将5舍去，5前为奇数应将5进位。
########################################
def round_liug(num, poi=0):
    num = float(num)
    num = num * pow(10, poi + 1)
    num_fir = num % 10
    num_sec = int(num / 10 % 10)
    if num_fir < 5 or num_fir > 5:
        if num_fir > 5:  # 大于五进位 小于五无视
            num += 10
    elif num_sec % 2 == 1:  # 保留数最后一位是奇数 进位，偶数无视
        num += 10
    num = float(math.trunc((num - num_fir) / 10))
    num = num / pow(10, poi)
    if poi == 0:
        num = round(num)
    else:
        ff = '.' + str(poi) + 'f'
        num = format(num, ff)
    return num


##########################################################
# 格式化字符串函数
# 输入为字符串，内容为包含正整数序列的列表（可能包含*字符），输出为字符串：
# 1、列表元素之前用'、'分隔；
# 2、对于大于等于3个以上的连续数字，合并为“最小值~最大值”的形式输出，例如：
# 输入：[1,3,4,5,6,13,14,19,29,30,31,32,34,40]
# 输出str：“1、3~5、6、13、14、19、29~34、40"
##########################################################
def format_lst(lst_str):
    lst = [int(s.replace('*', '')) for s in lst_str]
    ret = []
    start = end = lst[0]
    for num in lst[1:]:
        if num == end + 1:
            end = num
        else:
            ret.append((start, end))
            start = end = num
    ret.append((start, end))
    # 下述一行代码会将两个连续数字的序列表示为 “7~8”，而不是“7、8”，
    # ret = ','.join(str(start) if start == end else f'{start}~{end}' for (start, end) in ret)
    out = []
    for (s, e) in ret:
        if s == e:
            out.append(str(s))
        elif int(e) == int(s) + 1:
            out.append(str(s))
            out.append(str(e))
        else:
            out.append(f'{s}~{e}')
    ret = '、'.join(out)
    return ret


##########################################################
# 定义类中所用到的数据结构和常量
##########################################################
Area = namedtuple("Area", "min_row, max_row, min_col, max_col", defaults=(1, None, 1, None))
# 定义图片的宽度
WIDTH_REQ = Mm(54)
WIDTH_RESULT = Mm(20)
WIDTH_IMAGE = Mm(130)
WIDTH_ATTACH = Mm(155)
MAX_ROW = 100  # 当读取性能等EXEL文件时，由于返回的最大行数经常错误（maxrow=1），所以使用最大值

# 定义报告结论页中“检验依据”的最大标准数目。
# 如果超出这个数目，将会把“检验依据”另起一页。
MAX_CRITERIA = 10
# 一览表中最后的不适用项目说明的默认值：
NotSurport_comment = '被测设备不适用'

# 线程结束时发送到 log_queue
TASK_FINISH = "##任务完成##"

# error Code
CRITICAL_ERROR = 'critical'

# Word 中的换行符: \a   换页符：\f
# Report类继承自Thread对象，方便主模块将此作为线程使用
class Report(threading.Thread):
    '''
        # TestItem: {'part':,'type':,'row':,'num':,'name':,'subname':,'unit':,'require':,'result':,'comment':,'verdict':,'stub':}
        # part：取值范围为1或者2，1代表第一部分：网络信息安全，2代表第二部分：性能测试
        #################################################################
        # 只有标题中才会含有 stub 属性
        # stub = 0: 该标题不是末梢，它不含有任何测试项目
        # stub = 1: 该标题是末梢，但此次测试，包含的所有项目都未测
        # stub = 2: 该标题是末梢，含有已测试的项目
        #################################################################
        # type: 0: 网络信息安全和性能测试的级别，最高级别
        #       1: 一级标题
        #       2: 其他标题
        #       10: 非标题，说明文字等
        #       11: 测试项目类型1
        #       12: 测试项目类型2
        #       13: 测试项目类型3
        #################################################################
        # verdict: ['合格', '不合格', '--', 'ref']  含义： ‘--’ 为不支持项
        #################################################################
    '''

    def __init__(self, xlsm_file, task_type=2, is_revision_mode=False):
        super().__init__()
        self.daemon = True
        self._stop_event = threading.Event()
        self.xlsm_file = Path(xlsm_file)
        # 任务类型：
        # 0：原始记录
        # 1: 检验报告
        # 2：报告+记录
        self.task_type = task_type
        self.is_report = True
        self.is_revision_mode = is_revision_mode
        # 程序的执行目录
        if getattr(sys, 'frozen', False) and hasattr(sys, '_MEIPASS'):
            # log_show('running in a PyInstaller bundle')
            self.exe_dir = sys._MEIPASS
        else:
            # log_show('running in a normal Python process')
            self.exe_dir = Path.cwd()

        # 根据输入的参数准备其他要用到的变量：
        self.xlsm_dir = self.xlsm_file.parent
        if self.is_report:
            tpl_path = str(Path(self.exe_dir, r'templates', 'TestReport.docx'))
        else:
            tpl_path = str(Path(self.exe_dir, r'templates', 'TestRecord.docx'))
        self.tpl = DocxTemplate(tpl_path)
        self.template_dir = Path(tpl_path).parent
        self.workbook = xl.load_workbook(self.xlsm_file)

        # 其他暂时还无法赋值的参数：
        self.context = {}
        self.output_name = None  # PATH类型
        self.output_dir = None  # PATH类型
        # 2023年新增加的变量：
        self.test_items = []



    ##########################################################
    # 调用的主程序只需调用report对象中的此方法即可，将生成最终报告或word版原始记录，并返回生成的报告完整路径和名称
    ##########################################################
    def run(self):
        task_lst = []
        if self.task_type == 2:
            task_lst = [True, False]
        elif self.task_type == 1:
            task_lst = [True]
        else:
            task_lst = [False]

        # 判断最后的执行状态(结果为CRITICAL_ERROR的不是正常退出，其他都是）
        run_result = None
        try:
            while not self._stop_event.is_set():
                for state in task_lst:
                    self.is_report = state
                    if self.generate_report() == CRITICAL_ERROR:
                        run_result = CRITICAL_ERROR
                        self.stop()     # 发生错误时立即退出
                self.stop()     # 正常完成时退出
        except Exception as e:
            run_result = CRITICAL_ERROR
            log_show.critical(f"发生了严重错误：{e}")
            self.stop()     # 发生未被程序考虑的错误时立即退出

        # 向 log_queue 发送任务完成的信号，同时将生成报告的完整路径名传递给主线程
        if run_result == CRITICAL_ERROR:
            msg = f"{TASK_FINISH}{CRITICAL_ERROR}"
        else:
            msg = f"{TASK_FINISH}{self.output_name}"
        log_show.info(msg)

    ####################################################
    # 将所有任务串联起来，生成最终的报告或记录：
    ####################################################
    def generate_report(self):
        # 参数都已准备好，开始生成报告：
        name = "检验报告" if self.is_report else "原始记录"
        log_show.info('*' * 60)
        log_show.info(f"开始生成{name}")
        log_show.info('*' * 60)
        if self.is_report:
            task_names = [
                ['打开原始记录表格，并读取原始记录中的基本任务信息', 'generate_task_info'],
                ['插入报告中的检验依据', 'generate_criteria'],
                ['生成“检验结果”大表格', 'generate_result_table'],
                ['生成报告首页中的结论内容', 'generate_conclusion'],
                ['生成“检验结果一览表”的统计内容', 'generate_toc'],
                ['生成检验结果一览表中的不支持项目说明表', 'generate_notSupport'],
                ['生成“检验人员一览表”', 'generate_tester_tbl'],
                ['生成“检验用仪表”', 'generate_instrument'],
                ['插入附件中性能表格', 'generate_perform_tbl'],
                ['插入附件中的图片', 'generate_attach_images']

            ]
        else:
            task_names = [
                ['打开原始记录表格，并读取原始记录中的基本任务信息', 'generate_task_info'],
                ['插入Word中的检验依据', 'generate_criteria'],
                ['生成“检验结果”大表格', 'generate_result_table'],
                # ['生成报告首页中的结论内容', 'generate_conclusion'],
                # ['生成“检验结果一览表”的统计内容', 'generate_toc'],
                # ['生成检验结果一览表中的不支持项目说明表', 'generate_notSupport'],
                ['生成“检验人员一览表”', 'generate_tester_tbl'],
                ['生成“检验用仪表”', 'generate_instrument'],
                ['插入附件中性能表格', 'generate_perform_tbl'],
                ['插入附件中的图片', 'generate_attach_images']
            ]

        for i, task in enumerate(task_names):
            if hasattr(self, task[1]):
                # info = str(i + 1) + '、正在：' + task[0] + '...'
                log_show.info(task[0])
                if getattr(self, task[1])() == CRITICAL_ERROR:
                    log_show.critical(f"{task[0]} 任务失败了！！")
                    self.stop()
                    return CRITICAL_ERROR
                    # info = '已完成：' + task[0] + '。'
                    # log_show.info(info)
        # context 中的内容已经更新完毕，返回给调用函数进行word模板文件渲染即可
        log_show.info('开始根据文档模板进行最终结果的渲染')
        # autoescape默认值为False，渲染的文档中如果有 <"&'> 等字符会有问题。
        # autoescape=True 可以解决这一问题
        self.tpl.render(self.context, autoescape=True)
        try:
            self.tpl.save(str(self.output_name))
        except Exception as e:
            log_show.critical(f"保存文件失败！请确认下述文件是否已经打开：\n{self.output_name}")
            self.stop()
            return CRITICAL_ERROR

        # 如果还有附件文档，则在生成文档的基础上进行处理
        log_show.info('开始处理附加文档')
        self.generate_attach_document()

        # 删除word文档最后的空白页：
        log_show.info("开始查找并删除文档最后的空白页")
        self.remove_last_blank_page()

        # 更新Word文档域
        log_show.info('开始更新word文档中的域')
        self.update_word_fields()

        # 对于原始记录，打开文档的修订模式：
        if not self.is_report and self.is_revision_mode:
            self.set_docx_trackRevisions()
            log_show.info('打开原始记录的修订模式')

        log_show.info('*' * 60)
        log_show.info(f'渲染完成，{name}已生成！！')
        log_show.info("请双击左下角博鼎Logo快速打开目录查看")
        log_show.info('*' * 60)

        # 重命名原始记录
        # 获取不包含后缀的文件名：
        new_name = self.output_name.stem
        output_excel = Path(self.xlsm_file).parent / (new_name + '.xlsm')
        if str(output_excel) != str(self.xlsm_file):
            old_xlsm_file = self.xlsm_file
            try:
                self.xlsm_file = Path(self.xlsm_file).rename(output_excel)
            except Exception as e:
                self.xlsm_file = old_xlsm_file
                log_show.error(f'重命名原始记录表格文件时发生错误：{str(e)}')


    ##########################################################
    # 结束线程：  因任务时间太短，没有必要使用此函数来暂停任务
    ##########################################################
    def stop(self):
        self._stop_event.set()

    # get_**** 系列的方法为内部辅助使用的函数，辅助其他方法实现功能
    #   当dir_parent='images' 或 ‘data’ 时，返回的是原始记录下对应目录的文件
    #   当dir_parent='template' 及其他字符时，返回的是模板文件夹下‘images’目录下的文件
    def get_file(self, filename, dir_parent='images', makeCopy=True, name=''):
        if '\\' in filename or '/' in filename:  # filename包含路径信息
            file = Path(filename)
        else:
            if dir_parent == 'template':
                file = Path(self.template_dir, filename.strip())
            else:
                file = Path(self.xlsm_dir, dir_parent, filename.strip())

        if not file.exists():
            log_show.error(f"原始记录中找不到“{file}”文件!")
            return None

        # 将非模板路径下的文件拷贝到输出目录下的'others'文件夹：
        if str(Path(file).parent) != str(self.template_dir) and makeCopy and self.output_dir:
            if name:
                output_name = self.context['report_number'] + '_' + name + Path(file).suffix
            else:
                output_name = self.context['report_number'] + '_' + Path(file).name

            # 在输出目录中新建'others'文件夹，并将文件拷贝到此文件夹中：
            dest_dir = self.output_dir.joinpath('others')
            dest_dir.mkdir(parents=True, exist_ok=True)
            shutil.copy(file, Path(dest_dir.joinpath(output_name)))
        return file


    # 处理图片内容, tpl 为word 模板文件，file 为图片保存的文件名称
    def get_image(self, file, width=WIDTH_IMAGE, name=None):
        image = {}
        if Path(file).exists():
            if width:
                image['image'] = InlineImage(self.tpl, str(file), width=width)
            else:
                image['image'] = InlineImage(self.tpl, str(file))
            image['name'] = name if name else file.stem
        return image

    # 2023新增
    # 读取Excel文件中的特定区域，并返回列表数据：
    def get_excel_data(self, file, sheet, area):
        if not Path(file).exists():
            log_show.error(f"找不到'{file}'文件！")
            return None
        workbook = xl.load_workbook(file, read_only=True)
        if isinstance(sheet, int):  # 传入的sheet是数字，表示的是sheet页的索引Index
            active_sheet = workbook.worksheets[sheet]
        else:
            if sheet in workbook.sheetnames:  # 传入的sheet是sheet页的名称
                active_sheet = workbook[sheet]
            else:
                log_show.warning(f"找不到“{file}”文件的“{sheet}” sheet页！")
                return None

        # 去掉原有表格中的dimensions值，否则 max_row 和 max_column 返回的值可能不正确，特别是仪表和其他程序自动生成的表格经常会返回A1:A1
        # active_sheet.reset_dimensions()
        # active_sheet.calculate_dimension(force=True)
        if not area.max_row:
            area = area._replace(max_row=active_sheet.max_row)
        if not area.max_col:
            area = area._replace(max_col=active_sheet.max_column)

        # log_show(f"“{sheet}” sheet页中，行数为：{area.max_row}，列数为：{area.max_col}", 'blue')
        data = []
        for row_number, row in enumerate(active_sheet.iter_rows(*area, values_only=True), start=area.min_row):
            if any(row):
                row_data = ['' if i is None else str(i).strip() for i in row]  # 需要考虑单元格为数字0的情况，不能简单归为''
                # row_data.insert(0, row_number)
                row_data.append(row_number)
                data.append(row_data)
        workbook.close()
        return data

    # 处理检验结果sheet页中的数据，生成self.test_items
    def process_excel_data(self):  # 存放测试结果的字典结构：
        # 打开原始记录 ('templates/TestRecord.xlsx')
        max_col = 7
        area = Area(min_row=2, max_row=None, min_col=1, max_col=max_col)
        data = self.get_excel_data(self.xlsm_file, sheet='检验结果', area=area)
        if not data:
            log_show.error(f"不存在”检验结果“sheet页，或者”检验结果“sheet页中没有有效数据！！")
            return None

        keys = ['num', 'name', 'subname', 'unit', 'require', 'result', 'comment', 'row']
        lst = []  # 用于存储处理后的结果
        for row in data:
            lst.append(dict(zip(keys, row)))

        # 1: 判断ti的类型：
        i = 0  # 初始化指针
        while i < len(lst):
            if lst[i]['subname'] == '' and lst[i]['unit'] == '' and lst[i]['require'] == '':  # 标题或说明
                if '$' in lst[i]['num']:
                    lst[i]['num'] = str(lst[i]['num']).replace('$', '')
                    lst[i]['type'] = 10  # 项目编号中含有'$'字符，不是标题，而是说明之类文字
                else:
                    lst[i]['type'] = 2  # 为标题
                i += 1  # 将指针向后移动
            elif lst[i]['name'] and (not lst[i]['subname']):  # type = 11 or 12
                if i == len(lst) - 1 or (i < len(lst) - 1 and lst[i + 1]['name']):  # 最后一个元素
                    lst[i]['type'] = 11
                    i += 1  # 将指针向后移动
                else:
                    j = i + 1  # 初始化另一个指针
                    lst[i]['type'] = 12
                    while j < len(lst) and lst[j]['name'] == '' and lst[j]['subname'] == '':
                        lst[j]['type'] = 120
                        lst[j]['name'] = lst[i]['name']
                        j += 1  # 将指针向后移动
                    i = j  # 将指针移动到下一个区间的起始位置
            elif lst[i]['name'] and lst[i]['subname']:  # type=13
                j = i + 1  # 初始化另一个指针
                lst[i]['type'] = 13
                while j < len(lst) and lst[j]['name'] == '' and lst[j]['subname']:
                    lst[j]['type'] = 130
                    lst[j]['name'] = lst[i]['name']
                    j += 1  # 将指针向后移动
                i = j  # 将指针移动到下一个区间的起始位置
            else:  # 如果该元组不符合以上任何一种情况,数据有误
                lst[i]['type'] = 21
                log_show.error(f"原始记录中第 {lst[i]['row']} 行的数据有误，请检查修改！")
                i += 1  # 将指针向后移动

        # 2：
        # 判断检测要求和检验结果中是否有图片，以及图片路径是否合法
        # 增加 verdict 字段：
        for ti in lst:
            if ti['type'] < 11:
                ti['verdict'] = None
            else:  # 该行不是标题
                # 检查单位、标准要求及检验结果列是否有空白，如果有，提示后退出程序
                if not all([ti['unit'], ti['require'], ti['result']]):
                    log_show.error(f"请检查原始记录中第 {ti['row']} 行，数据可能不完整！")

                # 判断检测要求和检验结果中是否有图片，以及图片路径是否合法
                if '图片' in ti['require']:  # 检验要求中包含图片
                    file = ti['require'].split('图片')[-1].strip()  # 截取图片的文件名
                    file = self.get_file(file, 'template')
                    image = self.get_image(file, width=WIDTH_REQ)
                    if image:
                        ti['require'] = image['image']
                    else:
                        log_show.error(f"原始记录中第 {ti['row']} 行检测要求中的图片文件找不到！")

                if '图片' in ti['result']:  # 检验结果中包含图片
                    file = ti['result'].split('图片')[-1].strip()  # 截取图片的文件名
                    file = self.get_file(file)
                    image = self.get_image(file, width=WIDTH_RESULT)
                    if image:
                        ti['result'] = image['image']
                    else:
                        log_show.error(f"原始记录中第 {ti['row']} 行检验结果中的图片文件找不到！！")

                # # 根据序号（是否带*)判断是否为参考项目；根据 ’result‘ 填写 ’verdict‘
                # verdict: ['合格', '不合格', '--', 'ref']  含义： ‘--’ 为不支持项
                if ti['comment'] in ['不合格', 'F', 'Fail', 'Failed']:
                    ti['verdict'] = '不合格'
                elif ti['result'] in ['/', '--', '不支持', '不适用', '允许不支持']:
                    ti['result'] = '不适用'
                    ti['verdict'] = '--'
                elif '*' in ti['num']:
                    ti['verdict'] = 'ref'
                else:
                    ti['verdict'] = '合格'

        # 3: 增加 part 字段，判断标题层级：
        part = 1
        for ti in lst:
            num = ti['num']
            level = len(num.split('.')) if num.split('.')[-1] else (len(num.split('.')) - 1)
            if ti['type'] < 10:
                if '第一' in num:
                    ti['type'] = 0
                elif '第二' in num:
                    ti['type'] = 0
                    part = 2
                elif level < 2:  # 是一级标题
                    ti['type'] = 1
                ti['counter'] = Counter()  # 所有的标题项都加入了‘counter’键，防止后续读取时出现‘key error’
            ti['part'] = part

        # 4: 增加 stub 字段，对于 stub 标题增加统计项目，并将全部未测试的标题项目删除；
        # 测试项目中的num按照大排列重新编号
        i = 0
        seq = 1
        while i < len(lst):
            if i == len(lst) - 1:
                lst[i]['stub'] = 0
                self.test_items.append(lst[i])  # 最后一个元素直接加入结果列表
                i += 1
            elif lst[i]['type'] < 10 <= lst[i + 1]['type']:
                j = i + 2
                while j < len(lst) and lst[j]['type'] >= 10:
                    j += 1
                cnt = Counter([cc['verdict'] for cc in lst[(i + 1):j]])
                cnt['tested'] = cnt['合格'] + cnt['不合格'] + cnt['ref']  # 实测项目数 = 合格项目数 + 不合格项目数 + 参考项数
                cnt['total'] = cnt['tested'] + cnt['--']  # 应测项目数 = 实测项目数 + 不支持项目数
                if cnt['tested']:  # 有实际测试的项目时
                    lst[i]['counter'] = cnt
                    lst[i]['stub'] = 2
                    # 测试项目中的num按照大排列重新编号，参考项目序号前加“*”
                    ii = i + 1
                    while ii < j:
                        if lst[ii]['type'] > 10:
                            if lst[ii]['verdict'] == 'ref':
                                lst[ii]['num'] = '*' + str(seq)
                            else:
                                lst[ii]['num'] = str(seq)
                            seq += 1
                        ii += 1
                    self.test_items.extend(lst[i:j].copy())
                elif cnt['total']:  # 已测项目数为0，但应测项目数不为0的项目，stub赋值为1
                    lst[i]['counter'] = Counter()
                    lst[i]['stub'] = 1
                    self.test_items.append(lst[i])
                else:  # 后面都是 type=10 的注释项目：
                    lst[i]['stub'] = 0
                    self.test_items.extend(lst[i:j].copy())
                cnt = Counter()
                i = j
            else:
                lst[i]['stub'] = 0
                self.test_items.append(lst[i])  # 元素直接加入结果列表
                i += 1

        # 写入一级标题的统计数据：
        i = 0
        while i < len(self.test_items):
            if self.test_items[i]['type'] == 1:  # 一级标题
                j = i + 1
                while j < len(self.test_items) and self.test_items[j]['type'] != 1:
                    j += 1
                c1 = Counter([cc['verdict'] for cc in self.test_items[(i + 1):j] if cc['type'] > 10])
                c1['tested'] = c1['合格'] + c1['不合格'] + c1['ref']  # 实测项目数 = 合格项目数 + 不合格项目数 + 参考项数
                c1['total'] = c1['tested'] + c1['--'] if c1['tested'] else 0  # 应测项目数 = 实测项目数 + 不支持项目数
                self.test_items[i]['counter'] = c1
                i = j
            else:
                i += 1

        # for ti in self.test_items:
        #     if ti['type'] == 1:
        #         log_show(ti)

    # 读取TestCenter生成的性能表格（XLSX）的数据
    def get_performance(self, file_main, file_light=None):
        if not Path(file_main).exists():
            log_show.error(f"找不到'{file_main}'文件！")
            return None
        throughput = []
        latency = []
        frame_loss = []
        workbook = xl.load_workbook(file_main, read_only=True)
        # 判断仪表类型并准备数据:
        if 'Test Summary Table' in workbook.sheetnames:  # SPIRENT的仪表
            log_show.debug("使用了SPIRENT的仪表测试性能")
            area_throughput = Area(min_row=5, max_row=MAX_ROW, min_col=3, max_col=7)
            sheet_troughput = 'Test Summary Table'

            area_latency = Area(min_row=5, max_row=MAX_ROW, min_col=1, max_col=6)
            sheet_latency = 'Advanced Test Summary Ta'

            area_latency10 = Area(min_row=5, max_row=MAX_ROW, min_col=1, max_col=4)
            sheet_latency10 = 'Test Summary Table'

        elif '测试汇总表' in workbook.sheetnames:  # 信而泰的仪表
            log_show.debug("使用了信而泰的仪表测试性能")
            area_throughput = Area(min_row=5, max_row=MAX_ROW, min_col=1, max_col=5)
            sheet_troughput = '测试汇总表'

            area_latency = Area(min_row=5, max_row=MAX_ROW, min_col=1, max_col=6)
            sheet_latency = '高级测试汇总表'

            area_latency10 = Area(min_row=5, max_row=MAX_ROW, min_col=1, max_col=4)
            sheet_latency10 = '测试汇总表'

        else:
            log_show.error(f"找不到“{file_main}”文件的性能sheet页！")
            return None

        # 读取性能数据中的吞吐量值:
        log_show.info('读取性能数据中的吞吐量值')
        throughput_lst = []
        rows = self.get_excel_data(file_main, sheet=sheet_troughput, area=area_throughput)
        if rows:
            throughput = [[round_liug(row[0]), round_liug(row[4]), round_liug(row[3], 2)] for row in
                          rows]
            throughput_lst = [row[3] for row in rows]

        # # 读取性能数据中的吞吐量下时延和丢帧率:
        log_show.info('读取性能数据中的吞吐量下时延和丢帧率')
        rows = self.get_excel_data(file_main, sheet=sheet_latency, area=area_latency)
        if rows:
            log_show.info('读取性能数据中的线速丢帧率')
            frame_loss = [[round_liug(row[0]), round_liug(row[3], 3)] for row in rows if round_liug(row[1]) == 100]
            # 读取性能数据中的吞吐量下时延
            log_show.info('读取性能数据中的时延')
            # 读取吞吐量值在前面汇总的列表中、同时丢包率是0的数据：
            latency = [[round_liug(row[0]), round_liug(row[5], 2)] for row in rows if
                       row[2] in throughput_lst and round_liug(row[3], 3) == '0.000']
            # latency = [[round_liug(row[0]), round_liug(row[5], 2)] for row in rows if row[2] in throughput_lst and row[3] == '0']
        # log_show(latency)

        latency10 = None
        if Path(file_light).exists():
            # 读取轻载时延数据，通常在另外一个Excel文件file_light当中
            log_show.info('读取轻载时延数据')
            area = Area(min_row=5, max_row=MAX_ROW, min_col=1, max_col=4)
            rows = self.get_excel_data(file_light, sheet=sheet_latency10, area=area_latency10)
            if rows:
                latency10 = [[round_liug(row[0]), round_liug(row[3], 2)] for row in rows]
        # log_show(latency10)

        return throughput, latency, frame_loss, latency10

    # 生成规范化的文件名称 ：报告编号_厂家_设备名称_设备型号，并将文件名中的非法字符,用‘-’替换
    # 创建输出目录；生成输出的文件名称
    def set_formal_name(self):
        report_number = self.context['report_number']
        manufacturer = self.context['manufacturer']
        equipment_type = self.context['equipment_type']
        equipment_model = self.context['equipment_model']

        # 去除设备厂商名称中的“技术有限公司“、”技术公司”等信息以简化报告名称：
        removal = ["科技发展股份有限公司", "科技股份有限公司", "技术股份有限公司", "产业股份有限公司",
                   "科技有限责任公司", "科技有限公司", "技术有限公司", "股份有限公司", "有限责任公司", "有限公司"]
        for string in removal:
            manufacturer = manufacturer.replace(string, "")

        dst_dir = report_number + '_' + manufacturer + '_' + equipment_type + '_' + equipment_model
        if self.is_report:
            new_name = report_number + '_报告'
        else:
            new_name = report_number + '_记录'
        new_name = new_name + '_' + manufacturer + '_' + equipment_type + '_' + equipment_model
        # 去掉文件名中的非法字符,用‘-’替换
        dst_dir = re.sub(r'[/:*?"<>|+\\\s]', '-', dst_dir)
        new_name = re.sub(r'[/:*?"<>|+\\\s]', '-', new_name)

        self.output_dir = Path(self.xlsm_dir).joinpath(dst_dir)
        # 创建输出目录
        # parents：如果父目录不存在，是否创建父目录
        # exist_ok：只有在目录不存在时创建目录，目录已存在时不会抛出异常。
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.output_name = Path(self.output_dir).joinpath(new_name + '.docx')

    # ######################################################################
    # 以下函数为生成Word的具体内容：
    # generate_xxxx 函数为最终生成报告内容的方法,返回值如果为None，则为异常退出；正常退出返回True
    # 生成报告的基本任务信息
    def generate_task_info(self):
        if not Path(self.xlsm_file).exists():
            log_show.critical(f"找不到'{self.xlsm_file}' 原始表格文件！!")
            return CRITICAL_ERROR

        # 读取 “map” sheet页，获得变量的名称
        area = Area(min_row=2, max_row=None, min_col=3, max_col=4)
        rows = self.get_excel_data(self.xlsm_file, sheet='map', area=area)
        if not rows:
            log_show.critical("原始记录XLSM文档中未包含'map'页,获取模板中的变量名称失败！！")
            return CRITICAL_ERROR

        key_dic = {row[0]: row[1] for row in rows if row[1]}

        # 读取 “基本信息” sheet页，给变量赋值
        area = Area(min_row=2, max_row=None, min_col=3, max_col=4)
        rows = self.get_excel_data(self.xlsm_file, sheet='基本信息', area=area)
        if not rows:
            log_show.critical("原始记录XLSM文档中未包含'基本信息'页,获取任务基本信息失败！！")
            return CRITICAL_ERROR
        # info_dic = {key_dic[row[0]]: row[1] for row in rows if row[1]}
        info_dic = {}
        for row in rows:
            if row[0] and row[1]:
                info_dic[key_dic[row[0]]] = row[1]

            elif row[0] and not row[1]:
                log_show.critical(f"原始记录 “基本信息” sheet页中的 “{row[0]}” 还未填写，请确认！!")
                return CRITICAL_ERROR

        # 对于项目一览表中额外说明信息的处理。如果为“/" 或者 "无" 等字符长度小于2的内容，则处理为空字符。
        other_comment = info_dic.get('toc_other', '')
        if other_comment:
            if len(other_comment) < 2:
                info_dic['toc_other'] = ''

        # 设置分页符字符
        self.context['newpage'] = '\f'
        self.context.update(info_dic)

        # 生成输出的目录和输出的文件名称：
        self.set_formal_name()

        # 读取 “检验样品照片” sheet页，读取拍摄时间和拍摄地点：
        area = Area(min_row=2, max_row=3, min_col=2, max_col=3)
        rows = self.get_excel_data(self.xlsm_file, sheet='检验样品照片', area=area)
        if not rows:
            log_show.critical("原始记录XLSM文档中未包含'检验样品照片'sheet页！！")
            return CRITICAL_ERROR

        info_dic2 = {key_dic[row[0]]: row[1] for row in rows if row[1]}
        info_dic.update(info_dic2)

        # 读取 “检验样品照片” sheet页，获取图片的文件名：
        area = Area(min_row=6, max_row=None, min_col=2, max_col=4)
        rows = self.get_excel_data(self.xlsm_file, sheet='检验样品照片', area=area)
        num = 0
        image_lst = []
        for row in rows:
            hide = str(row[2]).strip() if row[2] else '否'
            if self.is_report and hide == '是':
                continue

            if not row[1]:
                continue

            image_file = self.get_file(str(row[1]))
            if image_file and image_file.exists():  # 读取图片文件名:
                # 获取文件名中去除文件后缀的内容：
                if row[0]:
                    name = str(row[0])
                else:
                    # 不带后缀名的文件名：
                    name = image_file.stem

                # image_dic 包含image['num']、image['title']、image['image']
                image_dic = self.get_image(image_file, width=WIDTH_IMAGE, name=name)

                if not image_dic:
                    continue
                num += 1
                image_dic['num'] = str(num)
                image_lst.append(image_dic.copy())

        info_dic['equipment_images'] = image_lst
        self.context.update(info_dic)

    # 生成报告中的“检验依据”
    def generate_criteria(self):
        # 读取 “检验依据” sheet页
        area = Area(min_row=2, max_row=None, min_col=1, max_col=3)
        rows = self.get_excel_data(self.xlsm_file, sheet='检验依据', area=area)
        if not rows:
            log_show.error(f"原始记录中找不到名称为 “检验依据” 的sheet页，请确认！")
            return CRITICAL_ERROR

        tbl_criteria = []
        str_criteria = ''
        ins = {}
        num = 0
        for row in rows:
            if row[2]:
                num += 1
                ins['num'] = str(num) + '.'
                ins['sn'] = row[1]
                ins['name'] = row[2]
                tbl_criteria.append(ins.copy())
                if num == 1:
                    str_criteria += f"{ins['num']}  {ins['sn']} {ins['name']}"
                else:
                    str_criteria += f"\a{ins['num']}  {ins['sn']} {ins['name']}"
        # log_show.info(str_criteria)
        self.context['tbl_criteria'] = tbl_criteria
        self.context['str_criteria'] = str_criteria

        # 读取‘基本信息’ sheet页中的 结论页报告依据中最大条数：
        maxCr = MAX_CRITERIA
        if '基本信息' in self.workbook.sheetnames:
            sheet1 = self.workbook['基本信息']
            maxGet = sheet1['D34'].value
            if maxGet:
                maxCr = int(maxGet)

        # 原始记录可以增加检验依据的数目 +4
        if not self.is_report:
            maxCr = maxCr + 4

        log_show.info(f"获取到结论页报告依据中最大条数为 {str(maxCr)}")
        if len(tbl_criteria) > maxCr:
            log_show.info("结论页标准依据数量超出阈值，已分页")
            self.context['has_two'] = True
            self.context['two_title'] = '\f检验依据：'
        else:
            self.context['has_two'] = False

    # 2023New：生成首页中的检验结论
    # verdict: ['合格', '不合格', '--', 'ref']  含义： ‘--’ 为不支持项
    # 实测项目数 = 合格项目数 + 不合格项目数 + 参考项数
    # 应测项目数 = 实测项目数 + 不支持项目数
    def generate_conclusion(self):
        con_ret = []
        for part in range(1, 3):
            c_part = Counter([t['verdict'] for t in self.test_items if t['type'] > 10 and t['part'] == part])
            # 应测项目数：
            n_total = c_part['合格'] + c_part['不合格'] + c_part['--'] + c_part['ref']
            # 实测项目数：
            n_tested = c_part['合格'] + c_part['不合格'] + c_part['ref']
            # 不支持项列表
            na_lst = [t['num'] for t in self.test_items if
                      t['type'] > 10 and t['part'] == part and t['verdict'] == '--']
            # 参考项列表
            ref_lst = [t['num'] for t in self.test_items if
                       t['type'] > 10 and t['part'] == part and t['verdict'] == 'ref']
            # 不合格项列表
            fail_lst = [t['num'] for t in self.test_items if
                        t['type'] > 10 and t['part'] == part and t['verdict'] == '不合格']

            str_na = '（第' + format_lst(na_lst) + '项）' if len(na_lst) else ''
            str_ref = '（第' + format_lst(ref_lst) + '项）' if len(ref_lst) else ''
            str_fail = '（第' + format_lst(fail_lst) + '项）' if len(fail_lst) else ''

            if n_total > 0:
                con_part = "应测项：根据被检设备情况及相应标准，共{0}项；\a".format(str(n_total))
                con_part += "允许不支持项：共{0}项{1}；\a".format(str(c_part['--']), str_na)
                con_part += "实测项：共{0}项，其中参考项{1}项{2}不做判定；\a".format(str(n_tested),
                                                                                  str(c_part['ref']), str_ref)
                con_part += "不合格项：共{0}项{1}；\a结论：合格".format(str(c_part['不合格']), str_fail)
            else:
                con_part = "应测项：根据被检设备情况及相应标准，共0项。"

            con_ret.append(con_part)

        self.context['conclusion1'] = con_ret[0]
        self.context['conclusion2'] = con_ret[1]

    # 2023New
    # 生成测试内容一览表
    def generate_toc(self):
        # ti['type'] = 0 ：目前只包含两个：第一部分：网络信息安全  、 第二部分：互联互通
        # ti['type'] = 1 ：1级标题
        # ti['stub'] = 0 ：非末梢标题
        # ti['stub'] = 1 ：末梢标题，但counter = {}
        # ti['stub'] = 2 ：末梢标题，但counter = 不为空
        def toc_assign(count):
            data['tested'] = count['tested']
            data['total'] = count['total'] if count['tested'] else 0
            data['notSupport'] = count['--'] if count['tested'] else '--'
            data['pass'] = count['合格'] if count['tested'] else '--'
            data['fail'] = count['不合格'] if count['tested'] else '--'
            data['ref'] = count['ref'] if count['tested'] else '--'
            return data

        tbl_toc = []
        data = {}
        c_total = Counter()
        for part in range(1, 3):
            tis = [ti for ti in self.test_items if ti['type'] < 10 and ti['part'] == part]
            if len(tis) == 0:
                break
            c_part = Counter()
            for ti in tis:
                if ti['stub'] == 0:  # 非末梢标题
                    data['type'] = 1
                else:  # 末梢标题
                    c_part += ti['counter']
                    data = toc_assign(ti['counter'])
                    data['type'] = 2
                data['num'] = ti['num']
                # 项目一览表中，如果标题内容中包含回车、中文的括号，则回车、中文括号中的内容都会被精简；而英文括号中的内容会被保留：
                data['title'] = ti['name'].split('\n')[0].split('（')[0].strip()

                tbl_toc.append(data.copy())
                data = {}
            # 添加合计：
            c_total += c_part
            data = toc_assign(c_part)
            data['type'] = 2
            data['num'] = '合计'
            data['title'] = ''
            # data['total'] = c_part['total']
            tbl_toc.append(data.copy())
            data = {}
        # 添加总合计：
        data = toc_assign(c_total)
        data['type'] = 2
        data['num'] = '共合计'
        data['title'] = ''
        # data['total'] = c_total['total']
        tbl_toc.append(data.copy())

        self.context['tbl_toc'] = tbl_toc

    # 2023New
    # 生成一览表中的允许不支持项情况说明列表：
    def generate_notSupport(self):
        tis = [ti for ti in self.test_items if ti['type'] > 10 and ti['verdict'] == '--']
        # tbl_lst = [{key: tis[key] for key in
        #             ['num', 'name', 'subname', 'comment']}]
        tbl_lst = []
        tbl_dic = {'num': '', 'name': '', 'comment': ''}

        if len(tis) == 0:
            return
        elif len(tis) == 1:
            tbl_dic['num'] = tis[0]['num']
            tbl_dic['comment'] = tis[0]['comment'] if tis[0]['comment'] else NotSurport_comment
            # 对于包含subname的不支持项目，也简单处理，只取name值
            tbl_dic['name'] = tis[0]['name']
            # 对于包含subname的不支持项目，使用'--'符合进行连接
            tbl_dic['name'] = tis[0]['name'] + '--' + tis[0]['subname'] if tis[0]['subname'] else tis[0]['name']

            # # 去除项目名称中的 “设备支持该功能时测试” 等信息以简化项目名称：
            # removal = ["(设备支持该功能时测试)", "（设备支持该功能时测试）", "(适用于支持该功能的设备)",
            #            "（适用于支持该功能的设备）"]
            # for string in removal:
            #     tbl_dic['name'] = tbl_dic['name'].replace(string, "").strip()

            tbl_lst.append(tbl_dic.copy())
        # 允许不支持列表长度大于等于2时：
        else:
            current_ti = tis[0]
            start_num = int(current_ti['num'].replace('*', ''))
            start_name = current_ti['name']
            start_subname = current_ti['subname']
            start_comment = current_ti['comment'] if current_ti['comment'] else NotSurport_comment
            end_num = start_num
            for i in range(1, len(tis)):
                next_ti = tis[i]
                next_num = int(next_ti['num'].replace('*', ''))
                next_name = next_ti['name']
                next_subname = next_ti['subname']
                next_comment = next_ti['comment'] if next_ti['comment'] else NotSurport_comment
                # 不适用项说明合并的三个条件：1、序列号连续  2、项目大类名称(name)一样  3、不适用说明的文字一样
                if next_num == end_num + 1 and next_name == start_name and next_comment == start_comment:
                    end_num = next_num
                else:
                    if start_num == end_num:
                        tbl_dic['num'] = str(start_num)
                        # 对于包含subname的不支持项目不合并时，使用'--'符合进行连接
                        tbl_dic['name'] = start_name + ' -- ' + start_subname if start_subname else start_name
                    else:
                        tbl_dic['num'] = f"{start_num}~{end_num}"
                        # 对于包含subname的不支持项目合并时，只使用大类名称（name）
                        tbl_dic['name'] = start_name
                    tbl_dic['comment'] = start_comment
                    tbl_lst.append(tbl_dic.copy())

                    start_num = next_num
                    start_name = next_name
                    start_subname = next_subname
                    start_comment = next_comment
                    end_num = start_num
            # 最后一个的处理：
            if start_num == end_num:
                tbl_dic['num'] = str(start_num)
                # 对于包含subname的不支持项目不合并时，使用'--'符合进行连接
                tbl_dic['name'] = start_name + ' -- ' + start_subname if start_subname else start_name
            else:
                tbl_dic['num'] = f"{start_num}~{end_num}"
                # 对于包含subname的不支持项目合并时，只使用大类名称（name）
                tbl_dic['name'] = start_name
            tbl_dic['comment'] = start_comment
            tbl_lst.append(tbl_dic.copy())

        # 去除不支持项目名称中的 “设备支持该功能时测试” 等信息以简化项目名称：
        for item in tbl_lst:
            removal = ["(设备支持该功能时测试)", "（设备支持该功能时测试）", "(适用于支持该功能的设备)",
                       "（适用于支持该功能的设备）"]
            for string in removal:
                item['name'] = item['name'].replace(string, "").strip()

        self.context['tbl_notSupport'] = tbl_lst

    # 2023新增加的测试代码：
    # 根据test_items的内容生成“检验结果”表格：
    def generate_result_table(self):
        # tbl_contens = [tbl1:{}, tbl2:{}]
        # tbl = {
        #           'type': ti.type,
        #           'num': ti.num,
        #           'title': ti.name,
        #           'data': [
        #                       ti_item1:{'no':ti.num,'name':, 'unit':, 'req','result', 'verdict'},
        #                       ti_item2:{'no':ti.num,'name':, 'unit':, 'req','result', 'verdict'},
        #                   ]
        #       }

        # 调用测试结果的预处理，生成 test_items 列表
        self.process_excel_data()

        # 初始化变量
        tbl_result = []
        tbl = {}
        data_lst = []
        i = 0
        while i < len(self.test_items):
            tbl['type'] = self.test_items[i]['type']
            tbl['title'] = self.test_items[i]['name']
            if tbl['type'] < 11:  # 标题或说明
                tbl['num'] = self.test_items[i]['num']
                tbl['data'] = []
                i += 1
            elif tbl['type'] == 11:
                dic_temp = {key: self.test_items[i][key] for key in
                            ['num', 'name', 'subname', 'unit', 'require', 'result', 'verdict', 'comment']}
                if dic_temp['verdict'] == 'ref':
                    dic_temp['verdict'] = '--'
                tbl['data'] = [dic_temp]
                i += 1
            elif tbl['type'] == 12 or tbl['type'] == 13:
                j = i + 1
                while j < len(self.test_items) and self.test_items[j]['type'] > 100:
                    j += 1
                ii = i
                while ii < j:
                    dic_temp = {key: self.test_items[ii][key] for key in
                                ['num', 'name', 'subname', 'unit', 'require', 'result', 'verdict', 'comment']}
                    if dic_temp['verdict'] == 'ref':
                        dic_temp['verdict'] = '--'
                    data_lst.append(dic_temp)
                    ii += 1
                tbl['data'] = data_lst.copy()
                data_lst = []
                i = j

            tbl_result.append(tbl.copy())
            tbl = {}

        # for ti in tbl_result:
        #     log_show(ti)
        self.context['tbl_result'] = tbl_result

    # 2023New:
    # 生成报告中的“检测仪表”表格
    def generate_instrument(self):
        # 读取 “检验用仪表” sheet页，获得检验用仪表列表
        area = Area(min_row=2, max_row=None, min_col=1, max_col=11)
        if '检验用仪表' not in self.workbook.sheetnames:
            log_show.error(f"原始记录中找不到名称为 “检验用仪表” 的sheet页，请确认！")
            return None
        rows = self.get_excel_data(self.xlsm_file, sheet='检验用仪表', area=area)
        tbl_instrument = []
        ins = []
        #  0        1       2       3       4       5       6       7       8        9
        # 序号    	仪表名称	型号	    生产厂家	出厂编号	硬件版本	软件版本	校准/验证 仪表来源  使用前/后状态
        #                                                           有效期至

        num = 0
        for row in rows:
            if row[9] and '正常' in row[9]:
                # log_show(row)
                num += 1
                ins.append(num)
                for i in range(1, 11):
                    ins.append(str(row[i]).strip() if row[i] else '--')
                tbl_instrument.append(ins.copy())
                ins = []
        self.context['tbl_instrument'] = tbl_instrument

    # 生成检测人员一览表--old
    def generate_tester_tbl_old(self):
        # ti.type = 0 ：目前只包含两个：第一部分：网络信息安全  、 第二部分：互联互通
        # ti.type = 1 ：1级标题
        tis = [ti for ti in self.test_items if ti['type'] < 2]
        tbl_lst = []
        tbl_dic = {'num': '', 'title': '', 'type': '', 'tester': '', 'auditor': ''}
        for ti in tis:
            tbl_dic['num'] = ti['num']
            # #取换行符、中文括号、英文括号前的部分
            tbl_dic['title'] = ti['name'].split('\n')[0].split('（')[0].split('(')[0].strip()
            if ti['type'] == 0:  # 目前只包含两个：第一部分：网络信息安全  、 第二部分：互联互通
                tbl_dic['type'] = 1
            else:
                tbl_dic['type'] = 2
                if ti['counter']['tested'] == 0:
                    tbl_dic['tester'] = '--'
                    tbl_dic['auditor'] = '--'
                else:
                    tbl_dic['tester'] = self.context['tester']
                    tbl_dic['auditor'] = self.context['auditor']
            tbl_lst.append(tbl_dic.copy())
        self.context['tbl_tester'] = tbl_lst

    # 2023 New：
    # 生成检测人员一览表
    def generate_tester_tbl(self):
        # 读取 “检验人员” sheet页，获取检验人员和检验时间
        area = Area(min_row=2, max_row=None, min_col=1, max_col=5)
        rows = self.get_excel_data(self.xlsm_file, sheet='检验人员', area=area)
        if not rows:
            log_show.error(f"原始记录中找不到名称为'检验人员'的sheet页，请确认！")
            return CRITICAL_ERROR

        #  0        1         2         3       4
        # 序号    	测试项目	  主检	    审核	    检验时间
        # ti.type = 1 ：目前只包含两个：第一部分：网络信息安全  、 第二部分：互联互通
        # ti.type = 2 ：1级标题
        tbl_lst = []
        tbl_dic = {'num': '', 'title': '', 'type': '', 'tester': '', 'auditor': '', 'date': ''}
        for row in rows:
            if row[0] and row[1]:  # 确保前两列不为空
                tbl_dic["num"] = row[0]
                tbl_dic["title"] = row[1]
                # tbl_dic["title"] = row[1].split('\n')[0].split('（')[0].split('(')[0].strip()
                if row[2] and row[3] and row[4]:
                    tbl_dic["type"] = 2
                    tbl_dic["tester"] = row[2]
                    tbl_dic["auditor"] = row[3]
                    tbl_dic["date"] = row[4]
                else:
                    tbl_dic["type"] = 1
                tbl_lst.append(tbl_dic.copy())
        self.context['tbl_tester'] = tbl_lst

    # 生成报告附件中的性能测试表格
    def generate_perform_tbl(self):
        # 读取 “附件” sheet页，获取性能数据的文件名：
        if '传输性能' not in self.workbook.sheetnames:
            log_show.warning(f"原始记录中找不到名称为 “传输性能” 的sheet页，请确认!")
            return None
        area = Area(min_row=3, max_row=7, min_col=2, max_col=5)
        rows = self.get_excel_data(self.xlsm_file, sheet='传输性能', area=area)
        num = 0
        perform = {}
        perform_lst = []
        for row in rows:
            if row[1]:  # 读取性能数据中的吞吐量值:
                filename = row[1].strip()
                file_main = self.get_file(filename, 'data')
                file_light = None
                if not file_main.exists():
                    log_show.warning(f"附件中的性能数据文件找不到:  {file_main.name}！")
                    continue
                num += 1
                if row[2]:  # 存在轻载时延数据文件
                    filename = row[2]
                    file_light = self.get_file(filename, 'data')
                    if not file_light.exists():
                        log_show.warning(f"附件中的轻载性能数据文件找不到:  {file_light.name}！")
                if num == 1:
                    perform['num'] = '表' + str(num) + ' '
                else:
                    perform['num'] = '\f\n表' + str(num) + ' '
                perform['title'] = row[0]
                perform['ports'] = row[3]
                perform['throughput'], perform['latency'], perform['frame_loss'], perform[
                    'latency10'] = self.get_performance(file_main, file_light)
                perform_lst.append(perform.copy())
        # log_show(perform_lst)
        self.context['perform_lst'] = perform_lst

    # 报告附件中插入普通图片（如眼图等）
    def generate_attach_images(self):
        # 读取 “附件” sheet页，获取图片的文件名：
        area = Area(min_row=3, max_row=20, min_col=2, max_col=5)
        if '附件' not in self.workbook.sheetnames:
            log_show.warning(f"原始记录中找不到名称为 “附件” 的sheet页。")
            return None
        rows = self.get_excel_data(self.xlsm_file, sheet='附件', area=area)
        num = 0
        image_lst = []
        for row in rows:
            hide = str(row[3]).strip() if row[3] else '否'
            if self.is_report and hide == '是':
                continue
            if row[2]:  # 读取图片文件名:
                # 获取文件名中去除文件后缀的内容：
                if row[1]:
                    name = str(row[1])
                else:
                    name = str(row[2])
                    last_dot_index = name.rfind('.')
                    if last_dot_index != -1:
                        name = name[:last_dot_index]

                # image_dic 包含image['num']、image['title']、image['image']
                image_file = self.get_file(str(row[2]))
                image_dic = self.get_image(image_file, width=WIDTH_IMAGE, name=name)
                if not image_dic:
                    continue
                num += 1
                image_dic['num'] = str(num)
                if row[0]:
                    image_dic['title'] = str(row[0]).strip()
                image_lst.append(image_dic.copy())
        self.context['attachment_images'] = image_lst

    # 附件中插入文档
    def generate_attach_document(self):
        # 读取 “附件” sheet页，获取文档的文件名：
        area = Area(min_row=23, max_row=30, min_col=2, max_col=5)
        rows = self.get_excel_data(self.xlsm_file, sheet='附件', area=area)
        if not rows:
            log_show.warning(f"原始记录中未找到需要附加的文档")
            return None

        middle_docx = None
        for row in rows:
            hide = str(row[3]).strip() if row[3] else '否'
            if self.is_report and hide == '是':
                continue
            if not row[2]:
                continue
            # 读取文件名:
            attach_file = self.get_file(str(row[2]), 'data')
            if attach_file and attach_file.exists():
                attach = Document(str(attach_file))
                if not middle_docx:  # 第一次时执行
                    master = Document(str(self.output_name))
                    # master.add_page_break()   # 文档之间加入分页符
                    middle_docx = Composer(master)
                # attach.add_page_break()
                middle_docx.append(attach)
        if middle_docx:
            middle_docx.save(str(self.output_name))
        return None

    # 打开word文档的修订模式：
    def set_docx_trackRevisions(self):
        # 打开文档
        doc = Document(str(self.output_name))
        # 在settings.xml中启用修订模式
        settings = doc.part.settings
        settings_xml = settings._element

        # 添加或更新<w:trackRevisions>标签
        track_revisions = settings_xml.find('w:trackRevisions', namespaces=settings_xml.nsmap)
        if track_revisions is None:
            # 创建新标签（启用修订）
            new_element = parse_xml(f'<w:trackRevisions {nsdecls("w")} w:val="1"/>')
            settings_xml.append(new_element)
        else:
            # 修改现有标签
            track_revisions.set('w:val', '1')  # '1'=启用, '0'=禁用

        # 保存文档
        doc.save(str(self.output_name))

    # 更新word文档域
    # 强制全选并更新（模拟Ctrl+A + F9）
    def update_word_fields(self):
        """
        使用win32com控制Word应用程序来更新域
        这种方法模拟Ctrl+A和F9的效果
        """
        word = None
        doc = None
        try:
            # 启动Word应用程序
            word = win32.Dispatch("Word.Application")
            word.Visible = False  # 隐藏Word窗口
            word.DisplayAlerts = False  # 禁用所有警告对话框

            # 打开文档
            doc = word.Documents.Open(str(self.output_name))
            time.sleep(2)
            
            # 方法1: 更新整个文档中的所有域
            doc.Fields.Update()
            
            # 方法2: 全选文档内容并更新域（备用方法）
            word.Selection.WholeStory()
            word.Selection.Fields.Update()

            # 保存文档
            doc.Save()

            log_show.info("Word文档域更新完成")

        except Exception as e:
            log_show.warning(f"Word文档域更新操作失败: {e}")

        finally:
            # 清理资源
            if doc:
                try:
                    doc.Close()
                except:
                    pass
            if word:
                try:
                    word.Quit()
                except:
                    pass

    # 删除Word文档中最后的空白页
    def remove_last_blank_page(self):
        try:
            # 打开Word文档
            doc = Document(str(self.output_name))

            # 获取文档中的所有段落
            paragraphs = doc.paragraphs

            # 从最后一个段落开始，向前查找空白段落并删除
            for i in range(len(paragraphs) - 1, -1, -1):
                # 检查段落是否为空白（包括只有空白字符的情况）
                if paragraphs[i].text.strip() == '':
                    p = paragraphs[i]._element
                    p.getparent().remove(p)

                else:
                    break  # 一旦遇到非空白段落，停止删除

            # 进一步检查文档末尾是否有分页符等导致的空白页
            # 获取文档最后一个元素
            last_element = doc.element.body[-1]

            # 如果最后一个元素是分页符，则删除
            if hasattr(last_element, 'tag') and 'sectPr' in last_element.tag:
                # 检查是否有连续的分页符
                for i in range(len(doc.element.body) - 2, -1, -1):
                    element = doc.element.body[i]
                    # 检查是否是分页符
                    if 'p' in element.tag and len(element) > 0:
                        pPr = element.find('.//w:pPr', namespaces=element.nsmap)
                        if pPr is not None:
                            pageBreakBefore = pPr.find('.//w:pageBreakBefore', namespaces=element.nsmap)
                            if pageBreakBefore is not None:
                                element.getparent().remove(element)
                            else:
                                break
                    else:
                        break

            # 保存修改后的文档
            doc.save(str(self.output_name))
            log_show.info(f"文档已查找并删除最后的空白页。")

        except Exception as e:
            log_show.error(f"删除空白页时出错：{str(e)}")




